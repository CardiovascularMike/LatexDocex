from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import matplotlib.pyplot as plt
from PIL import Image
from docx.oxml import OxmlElement

class DocFormatter:
    GreekDictionary = {
        '\\alpha': 'α',
        '\\beta': 'β',
        '\\gamma': 'γ',
        '\\delta': 'δ',
        '\\epsilon': 'ε',
        '\\zeta': 'ζ',
        '\\eta': 'η',
        '\\theta': 'θ',
        '\\iota': 'ι',
        '\\kappa': 'κ',
        '\\lambda': 'λ',
        '\\mu': 'μ',
        '\\nu': 'ν',
        '\\xi': 'ξ',
        '\\omicron': 'ο',
        '\\pi': 'π',
        '\\rho': 'ρ',
        '\\sigma': 'σ',
        '\\tau': 'τ',
        '\\upsilon': 'υ',
        '\\phi': 'φ',
        '\\chi': 'χ',
        '\\psi': 'ψ',
        '\\omega': 'ω',
        '\\cdot': '·',
        '\\times': '×',
        '\\approx': '≈',
        '\\%': '%'
    }

    def __init__(self):
        self.document = Document()

    def set_line_spacing(self, spacing=1.5):
        """Sets the line spacing for the Word document."""
        for paragraph in self.document.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = Pt(spacing * 12)  # 1.5 line spacing

    def add_paragraph(self, text):
        """Adds a paragraph to the Word document."""
        self.document.add_paragraph(text)

    def add_heading(self, text, level=1):
        """Adds a heading to the Word document."""
        self.document.add_heading(text, level=level)
        print('heading')

    def add_bold(self, line, bold_text):
        """Adds a bold text paragraph to the Word document."""
        print('DEBUG: Bold Line:', line)
        print('DEBUG: Bold Text:', bold_text)
        para = self.document.add_paragraph()
        lower_line = line.lower()
        lower_bold_text = bold_text.lower()
        start_pos = lower_line.find(lower_bold_text)
        end_pos = start_pos + len(bold_text)
        pre_bold_text = line[:start_pos]
        post_bold_text = line[end_pos:]

        if pre_bold_text:
            para.add_run(pre_bold_text)
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        if post_bold_text:
            para.add_run(post_bold_text)
        
    def add_bulleted_line(self, text):
        """Adds a bulleted list item to the Word document."""
        """Delete the first character of the text, which is the *."""
        text = text.strip()[1:]
        self.document.add_paragraph(text, style='List Bullet')
        print('bulleted line')

    def add_figure(self, filename):
        """Adds a figure to the Word document."""
        self.document.add_picture('PlaceImagesHere/'+filename, width=Inches(6.0))
        print('figure')

    def add_caption(self, text, figureNumber):
        """Adds a caption to the Word document."""
        para = self.document.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run('Figure ' + str(figureNumber) + ': ' + text)
        run.italic = True

    def add_math_paragraph(self, parts):
        """Adds a paragraph with inline math to the Word document."""
        para = self.document.add_paragraph()
        for text, math_text in parts:
            if text:
                para.add_run(text)
            if math_text:
                # Check for \mbox and ignore the command, just use the text part
                math_text = re.sub(r'\\mbox\{(.*?)\}', r'\1 ', math_text)

                # Check for the presence of \sqrt, \sum, or \frac
                if any(cmd in math_text for cmd in ['\\sqrt', '\\sum', '\\frac', '\\tfrac']):
                    # Send the entire math text to matplotlib as a plt.text statement
                    math_text = re.sub(r'\\tfrac', r'\\frac', math_text)
                    self.add_math_as_figure(math_text, para)
                else:
                    #-Replace greek LaTeX commands with greek characters-#
                    for latex_code, greek_char in self.GreekDictionary.items():
                        math_text = math_text.replace(latex_code, greek_char)
                    # Handle superscripts and subscripts
                    #-Look for carrot with bracketed stuff or with a single character afterwards-#
                    #-Then look for subscript with bracketed stuff or signle character-#
                    exp_pattern = re.compile(r'\^(\{.*?\}|\S)')
                    sub_pattern = re.compile(r'_(\{.*?\}|\S)')
                    last_end = 0
                    segments = []
                    for match in re.finditer(r'(\^|\_)(\{.*?\}|\S)', math_text):
                        base_text = math_text[last_end:match.start()]
                        segments.append((base_text, {}))
                        symbol = match.group(1)
                        formatted_text = match.group(2).strip('{}')
                        if symbol == '^':
                            segments.append((formatted_text, {'superscript': True}))
                        elif symbol == '_':
                            segments.append((formatted_text, {'subscript': True}))
                        last_end = match.end()
                    remaining_text = math_text[last_end:]
                    segments.append((remaining_text, {}))

                    for segment_text, formatting in segments:
                        run = para.add_run(segment_text)
                        run.font.italic = True
                        if 'superscript' in formatting:
                            run.font.superscript = True
                        if 'subscript' in formatting:
                            run.font.subscript = True

        print('math paragraph')

    def add_math_as_figure(self, math_text, para=False):
        """Adds a math expression as a figure to the Word document using matplotlib."""
        if not para:
            print('Not INLINE MATH FOUND!DEBUG!')
            para = self.document.add_paragraph()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # This is a LaTeX math expression, remove the $ symbols]
        self.add_math_latex_as_figure(math_text, para)

    def add_math_latex_as_figure(self, math_text, para):
        """Helper method to add a LaTeX math expression as a figure."""
        fig, ax = plt.subplots()
        ax.axis('off')
        text = ax.text(0.5, 0.5, f'${math_text}$', fontsize=10, ha='center', va='center')

        # Draw the canvas to get the bounding box of the text
        fig.canvas.draw()
        bbox = text.get_window_extent()
        bbox = bbox.transformed(fig.dpi_scale_trans.inverted())
        fig_width, fig_height = bbox.width, bbox.height
        fig.set_size_inches(fig_width, fig_height)
        fig.canvas.draw()
        
        fig.savefig('math_expression.png', dpi=192, bbox_inches='tight')
        image = Image.open('math_expression.png')
        width, height = image.size

        dpi_x, dpi_y = image.info.get('dpi')
        crop_fraction = 0.7
        crop_height = int(height * crop_fraction)
        crop_top = (height - crop_height) // 3
        crop_bottom = crop_top + crop_height
        cropped_image = image.crop((0, crop_top, width, crop_bottom))

        new_width1, new_height1 = image.size
        new_width = int(new_width1 * 0.8)
        new_height = int(new_height1 * 0.8)
        resized_image = image.resize((new_width, new_height))
        cropped_image.save('math_expression2.png')
        inchwidth = new_width / dpi_x
        # Load the saved image and add it to the document
        run = para.add_run()
        run.add_picture('math_expression2.png', width=Inches(inchwidth))
        print('math figure added')

    def save(self, filename):
        """Saves the Word document to the specified file."""
        self.document.save(filename)
