from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

class DocFormatter:
    def __init__(self):
        self.document = Document()

    def add_paragraph(self, text):
        """Adds a paragraph to the Word document."""
        self.document.add_paragraph(text)

    def add_heading(self, text, level=1):
        """Adds a heading to the Word document."""
        self.document.add_heading(text, level=level)
        print('heading')

    def add_bold(self, text):
        """Adds a bold text paragraph to the Word document."""
        para = self.document.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        print('bold')

    def add_bulleted_line(self, text):
        """Adds a bulleted list item to the Word document."""
        """Delete the first character of the text, which is the *."""
        print(text)
        text = text.strip()[1:]
        print(text)
        self.document.add_paragraph(text, style='List Bullet')
        print('bulleted line')

    def add_figure(self, filename):
        """Adds a figure to the Word document."""
        self.document.add_picture('PlaceImagesHere/'+filename, width=Inches(6.0))
        print('figure')

    def add_caption(self, text, figureNumber):
        """Adds a caption to the Word document."""
        para = self.document.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run('Figure ' + str(figureNumber) + ': ' + text)
        run.italic = True

    def add_math_paragraph(self, parts):
        """Adds a paragraph with inline math to the Word document."""
        para = self.document.add_paragraph()
        exp_pattern = re.compile(r'\^(\{.*?\}|\S)')
        sub_pattern = re.compile(r'_(\{.*?\}|\S)')
        sum_pattern = re.compile(r'\\sum')
        mbox_pattern = re.compile(r'\\mbox\{(.*?)\}')
        sqrt_pattern = re.compile(r'\\sqrt\{(.*?)\}')
        
        for text, math_text in parts:
            if text:
                para.add_run(text)
            if math_text:
                last_end = 0
                segments = []
                for match in re.finditer(r'(\^|\_)(\{.*?\}|\S)|\\sum|\\mbox\{(.*?)\}|\\sqrt\{(.*?)\}', math_text):
                    base_text = math_text[last_end:match.start()]
                    segments.append((base_text, {}))
                    if match.group(0) == '\\sum':
                        segments.append(('∑', {}))
                    elif match.group(0).startswith('\\mbox'):
                        segments.append((match.group(3), {}))
                    elif match.group(0).startswith('\\sqrt'):
                        sqrt_content = match.group(4)
                        segments.append(('√', {}))
                        segments.append((f"({sqrt_content})", {}))  # Fallback to parentheses
                    else:
                        symbol = match.group(1)
                        formatted_text = match.group(2).strip('{}')
                        if symbol == '^':
                            segments.append((formatted_text, {'superscript': True}))
                        elif symbol == '_':
                            segments.append((formatted_text, {'subscript': True}))
                    last_end = match.end()
                remaining_text = math_text[last_end:]
                segments.append((remaining_text, {}))

                for segment_text, formatting in segments:
                    run = para.add_run(segment_text)
                    run.font.italic = True
                    if 'superscript' in formatting:
                        run.font.superscript = True
                    if 'subscript' in formatting:
                        run.font.subscript = True

        print('math paragraph')

    def save(self, filename):
        """Saves the Word document to the specified file."""
        self.document.save(filename)
